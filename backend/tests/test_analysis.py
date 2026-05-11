"""Unit tests for app.services.analysis."""
import io
import os
import uuid
from unittest.mock import MagicMock, patch

import pytest

from app.services.analysis import (
    _detect_audio_extension,
    _extract_text_from_raw,
    _mock_analysis,
    analyse_submission,
)
from app.models.db import Submission, SubmissionFile, AnalysisResult


# ---------------------------------------------------------------------------
# _detect_audio_extension
# ---------------------------------------------------------------------------

class TestDetectAudioExtension:
    def test_mp3_id3(self):
        assert _detect_audio_extension(b"ID3" + b"\x00" * 10) == ".mp3"

    def test_mp3_frame_sync(self):
        # 0xFF 0xFB is a common MP3 frame sync
        assert _detect_audio_extension(b"\xFF\xFB" + b"\x00" * 10) == ".mp3"

    def test_wav(self):
        assert _detect_audio_extension(b"RIFF" + b"\x00" * 10) == ".wav"

    def test_flac(self):
        assert _detect_audio_extension(b"fLaC" + b"\x00" * 10) == ".flac"

    def test_ogg(self):
        assert _detect_audio_extension(b"OggS" + b"\x00" * 10) == ".ogg"

    def test_mp4(self):
        data = b"\x00\x00\x00\x08" + b"ftyp" + b"\x00" * 10
        assert _detect_audio_extension(data) == ".mp4"

    def test_webm(self):
        assert _detect_audio_extension(b"\x1a\x45\xdf\xa3" + b"\x00" * 10) == ".webm"

    def test_unknown_defaults_to_mp3(self):
        assert _detect_audio_extension(b"\x00\x01\x02\x03") == ".mp3"

    def test_empty_bytes_defaults_to_mp3(self):
        assert _detect_audio_extension(b"") == ".mp3"


# ---------------------------------------------------------------------------
# _extract_text_from_raw
# ---------------------------------------------------------------------------

class TestExtractTextFromRaw:
    def test_plain_utf8(self):
        text = _extract_text_from_raw(b"hello world", "document")
        assert text == "hello world"

    def test_utf8_with_special_chars(self):
        data = "café résumé".encode("utf-8")
        assert "café" in _extract_text_from_raw(data, "document")

    def test_invalid_utf8_replaced(self):
        data = b"valid\xff\xfe invalid"
        result = _extract_text_from_raw(data, "document")
        assert "valid" in result  # replacement mode, no exception

    def test_docx_extraction(self):
        """Build a minimal in-memory .docx and verify text extraction."""
        import docx as _docx
        buf = io.BytesIO()
        doc = _docx.Document()
        doc.add_paragraph("Top secret evidence paragraph.")
        doc.save(buf)
        raw = buf.getvalue()
        result = _extract_text_from_raw(raw, "document")
        assert "Top secret evidence paragraph." in result

    def test_docx_table_cells_extracted(self):
        """Table cell content must be included."""
        import docx as _docx
        buf = io.BytesIO()
        doc = _docx.Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "CellA"
        table.cell(0, 1).text = "CellB"
        doc.save(buf)
        raw = buf.getvalue()
        result = _extract_text_from_raw(raw, "document")
        assert "CellA" in result
        assert "CellB" in result

    def test_xlsx_extraction(self):
        """Build a minimal in-memory .xlsx and verify text extraction."""
        import openpyxl
        buf = io.BytesIO()
        wb = openpyxl.Workbook()
        ws = wb.active
        ws["A1"] = "Revenue"
        ws["B1"] = 999999
        wb.save(buf)
        raw = buf.getvalue()
        result = _extract_text_from_raw(raw, "spreadsheet")
        assert "Revenue" in result
        assert "999999" in result

    def test_binary_office_no_content_returns_stub(self):
        """A valid ZIP that is neither docx nor xlsx returns a fallback stub."""
        import zipfile
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            z.writestr("nothing.txt", "irrelevant")
        raw = buf.getvalue()
        result = _extract_text_from_raw(raw, "document")
        # Should return stub text, not raise
        assert isinstance(result, str)

    def test_csv_plain_text(self):
        data = b"col1,col2\nval1,val2\n"
        result = _extract_text_from_raw(data, "spreadsheet")
        assert "col1" in result


# ---------------------------------------------------------------------------
# _mock_analysis
# ---------------------------------------------------------------------------

class TestMockAnalysis:
    def test_returns_high_reliability(self):
        result = _mock_analysis(["document", "spreadsheet"])
        assert result.reliability_class == "HIGH"

    def test_corroborating_sources_capped_at_3(self):
        result = _mock_analysis(["a", "b", "c", "d", "e"])
        assert result.corroborating_sources <= 3

    def test_evidence_types_preserved(self):
        types = ["document", "image"]
        result = _mock_analysis(types)
        assert result.evidence_types == types

    def test_attribution_sentences_present(self):
        result = _mock_analysis(["document"])
        assert len(result.attribution_sentences) >= 1

    def test_scores_in_range(self):
        result = _mock_analysis(["document"])
        for score in [result.consistency_score, result.corroboration_score,
                      result.plausibility_score, result.overall_confidence]:
            assert 0 <= score <= 100


# ---------------------------------------------------------------------------
# analyse_submission — integration with mock analysis
# ---------------------------------------------------------------------------

def _make_submission_with_file(db, tmp_path) -> str:
    """Create a minimal Submission + SubmissionFile + encrypted file."""
    from app.core.crypto import encrypt_bytes, sha256_bytes, derive_file_key

    master_key = os.urandom(32)
    content = b"test evidence content"
    content_hash = sha256_bytes(content)
    submission_id = uuid.uuid4()
    file_key = derive_file_key(master_key, str(submission_id), content_hash)
    enc = encrypt_bytes(content, file_key)

    enc_path = tmp_path / "file.enc"
    enc_path.write_bytes(enc)

    submission = Submission(
        id=submission_id,
        merkle_root="a" * 64,
        file_count=1,
        status="TIMESTAMPED",
    )
    db.add(submission)
    db.flush()

    sf = SubmissionFile(
        id=uuid.uuid4(),
        submission_id=submission_id,
        filename_hash=sha256_bytes(b"filename.txt"),
        content_hash=content_hash,
        file_type="document",
        size_bytes=len(content),
        merkle_proof_path=[],
        encrypted_path=str(enc_path),
    )
    db.add(sf)
    db.flush()

    return str(submission_id), master_key


class TestAnalyseSubmission:
    def test_creates_analysis_result(self, db, tmp_path):
        submission_id, master_key = _make_submission_with_file(db, tmp_path)

        with patch("app.services.analysis.settings") as mock_settings:
            mock_settings.openai_api_key = ""  # triggers mock path
            mock_settings.encryption_key = master_key
            mock_settings.openai_max_chars_total = 400_000

            analyse_submission(submission_id, db)

        result = db.query(AnalysisResult).filter(
            AnalysisResult.submission_id == uuid.UUID(submission_id)
        ).first()
        assert result is not None
        assert result.overall_confidence == 75

    def test_updates_submission_status_to_analyzed(self, db, tmp_path):
        submission_id, master_key = _make_submission_with_file(db, tmp_path)

        with patch("app.services.analysis.settings") as mock_settings:
            mock_settings.openai_api_key = ""
            mock_settings.encryption_key = master_key
            mock_settings.openai_max_chars_total = 400_000

            analyse_submission(submission_id, db)

        submission = db.query(Submission).filter(
            Submission.id == uuid.UUID(submission_id)
        ).first()
        assert submission.status == "ANALYZED"

    def test_raises_for_missing_submission(self, db):
        with pytest.raises(ValueError, match="not found"):
            analyse_submission(str(uuid.uuid4()), db)

    def test_attribution_quotes_persisted(self, db, tmp_path):
        submission_id, master_key = _make_submission_with_file(db, tmp_path)

        with patch("app.services.analysis.settings") as mock_settings:
            mock_settings.openai_api_key = ""
            mock_settings.encryption_key = master_key
            mock_settings.openai_max_chars_total = 400_000

            analyse_submission(submission_id, db)

        result = db.query(AnalysisResult).filter(
            AnalysisResult.submission_id == uuid.UUID(submission_id)
        ).first()
        assert isinstance(result.attribution_quotes, list)
        assert len(result.attribution_quotes) >= 1
